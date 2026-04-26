import json
from pathlib import Path
from typing import Dict, Any, List

from .config import (
    PROJECT_TITLE,
    DATASET_NAME,
    DATASET_SOURCE_PAGE,
    DATASET_DIRECT_URL,
    DATASET_CITATION,
    REPORT_DIR,
    OUTPUTS_DIR,
    FIGURES_DIR,
    ANALYSIS_RESULTS_FILE,
    CLEANING_LOG_FILE,
)


def _fmt_money(value) -> str:
    if value is None:
        return "not available"
    return f"£{float(value):,.2f}"


def _fmt_number(value) -> str:
    if value is None:
        return "not available"
    return f"{float(value):,.0f}"


def _fmt_percent(value) -> str:
    if value is None:
        return "not available"
    return f"{float(value):.2f}%"


def _fmt_decimal(value) -> str:
    if value is None:
        return "not available"
    return f"{float(value):.3f}"


def _load_json(path: Path) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as file:
        return json.load(file)


def generate_markdown_report(student_name: str, student_id: str) -> Path:
    """Generate a submission-ready Markdown report using computed project results."""
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    results = _load_json(ANALYSIS_RESULTS_FILE)
    cleaning_log = _load_json(CLEANING_LOG_FILE)

    metrics = results["key_metrics"]

    report_path = REPORT_DIR / "Final_Project_Report.md"

    figure_references = [
        ("Figure 1", "Monthly Revenue Trend", "../outputs/figures/01_monthly_revenue_trend.png"),
        ("Figure 2", "Top 10 Countries by Revenue", "../outputs/figures/02_top_countries_by_revenue.png"),
        ("Figure 3", "Top 10 Products by Revenue", "../outputs/figures/03_top_products_by_revenue.png"),
        ("Figure 4", "Distribution of Invoice Values", "../outputs/figures/04_invoice_value_distribution.png"),
        ("Figure 5", "Relationship Between Distinct Products and Invoice Value", "../outputs/figures/05_invoice_items_vs_value_relationship.png"),
        ("Figure 6", "Invoice Value Comparison: UK vs International", "../outputs/figures/06_market_group_boxplot.png"),
    ]

    lines: List[str] = []

    lines.extend(
        [
            f"# {PROJECT_TITLE}",
            "",
            f"**Student name:** {student_name}",
            "",
            f"**Student ID:** {student_id}",
            "",
            f"**Dataset:** {DATASET_NAME}",
            "",
            f"**Dataset source page:** {DATASET_SOURCE_PAGE}",
            "",
            f"**Direct dataset file:** {DATASET_DIRECT_URL}",
            "",
            f"**Citation:** {DATASET_CITATION}",
            "",
            "---",
            "",
            "## 1. Project objective",
            "",
            "The objective of this project is to analyze customer purchase behavior and revenue trends using a real online retail transaction dataset. The analysis focuses on how revenue changes over time, which countries and products contribute most to sales, how customers purchase, and which transactions are unusually large.",
            "",
            "## 2. Analytical questions",
            "",
            "1. How does online retail revenue change over time?",
            "2. Which countries and products contribute most to total revenue?",
            "3. What patterns exist in customer purchase behavior, and which transactions look unusually large?",
            "",
            "These questions matter because an online retailer can use them to understand seasonality, identify important markets and products, improve customer strategy, and detect unusually large purchases that may need special attention.",
            "",
            "## 3. Dataset description and understanding",
            "",
            f"The raw dataset contains **{_fmt_number(metrics['raw_rows'])} rows** and **{_fmt_number(metrics['raw_columns'])} columns**. The analysis-ready sales dataset contains **{_fmt_number(metrics['analysis_ready_rows'])} rows** after cleaning. The transaction period in the cleaned data is from **{metrics['date_min']}** to **{metrics['date_max']}**.",
            "",
            "The dataset includes invoice number, product code, product description, quantity, invoice date, unit price, customer ID, and customer country. This structure is suitable for revenue analysis, product ranking, customer-level summaries, country comparison, time trend analysis, and outlier detection.",
            "",
            "Important data-understanding outputs were exported into the `outputs/tables/` folder:",
            "",
            "- `data_types_and_missing_values.csv`",
            "- `sample_rows.csv`",
            "- `numeric_summary_statistics.csv`",
            "",
            "## 4. Data cleaning and preparation",
            "",
            "The following cleaning steps were applied:",
            "",
        ]
    )

    for item in cleaning_log["cleaning_justification"]:
        lines.append(f"- **{item['step']}:** {item['reason']}")

    lines.extend(
        [
            "",
            "Cleaning summary:",
            "",
            f"- Duplicate rows removed: **{_fmt_number(cleaning_log['duplicate_rows_removed'])}**",
            f"- Cancelled or returned rows identified: **{_fmt_number(cleaning_log['cancelled_or_return_rows_identified'])}**",
            f"- Rows removed from positive-sales analysis: **{_fmt_number(cleaning_log['rows_removed_from_sales_analysis'])}**",
            f"- Missing customer ID rows remaining in sales data: **{_fmt_number(cleaning_log['missing_customer_id_rows_in_sales_data'])}**",
            "",
            "The project focuses on positive completed sales, so cancellations, returns, non-positive quantities, and non-positive prices were removed from the main revenue analysis. Missing customer IDs were not used for customer-level summaries because unidentified customers cannot be reliably grouped as individual buyers.",
            "",
            "## 5. Feature engineering",
            "",
            "Several derived columns were created to support analysis:",
            "",
            "- **total_revenue:** `quantity × unit_price`, used as the main sales value measure.",
            "- **invoice_month:** extracted from invoice date for monthly trend analysis.",
            "- **invoice_day_name and invoice_hour:** used for time-based behavior analysis.",
            "- **market_group:** separates United Kingdom transactions from international transactions.",
            "- **day_type:** separates weekday and weekend purchases.",
            "- **basket_value:** total value of each invoice.",
            "- **basket_distinct_items:** number of different products in each invoice.",
            "- **basket_value_category:** low, medium, or high basket value using NumPy percentile cutoffs.",
            "- **line_revenue_zscore:** NumPy-based standardized score for line revenue.",
            "",
            "These features make it possible to compare groups, analyze trends, examine customer behavior, and detect unusual high-value transactions.",
            "",
            "## 6. Analysis and visualizations",
            "",
            f"Overall, the cleaned data contains **{_fmt_number(metrics['total_orders'])} unique invoices**, **{_fmt_number(metrics['unique_customers'])} known customers**, **{_fmt_number(metrics['unique_products'])} products**, and **{_fmt_number(metrics['unique_countries'])} countries**. Total positive sales revenue is **{_fmt_money(metrics['total_revenue'])}**, with an average order value of **{_fmt_money(metrics['average_order_value'])}**.",
            "",
            "### 6.1 Monthly revenue trend",
            "",
            f"The highest revenue month is **{metrics['best_month']}**, with revenue of **{_fmt_money(metrics['best_month_revenue'])}**.",
            "",
            "![Monthly Revenue Trend](../outputs/figures/01_monthly_revenue_trend.png)",
            "",
            "This chart shows how revenue changes month by month. It helps identify seasonal changes and periods of stronger sales. The trend is important because online retailers often need to plan stock, marketing, and staffing around high-revenue months.",
            "",
            "### 6.2 Country-level revenue contribution",
            "",
            f"The top revenue country is **{metrics['top_country']}**, contributing **{_fmt_money(metrics['top_country_revenue'])}**, or **{_fmt_percent(metrics['top_country_share_percent'])}** of total revenue.",
            "",
            "![Top Countries by Revenue](../outputs/figures/02_top_countries_by_revenue.png)",
            "",
            "This chart compares revenue by country. It shows whether the business depends heavily on one market or has a balanced international customer base. This is a required subgroup comparison because it compares country groups using revenue, orders, and average order value.",
            "",
            "### 6.3 Product-level revenue contribution",
            "",
            f"The highest revenue product is **{metrics['top_product']}**, with revenue of **{_fmt_money(metrics['top_product_revenue'])}**.",
            "",
            "![Top Products by Revenue](../outputs/figures/03_top_products_by_revenue.png)",
            "",
            "This chart identifies products that generate the most revenue. This is useful for inventory planning because products with high revenue may deserve priority in stock management and promotion.",
            "",
            "### 6.4 Invoice value distribution and outliers",
            "",
            f"The 99th percentile invoice value is **{_fmt_money(metrics['invoice_value_p99'])}**. Using the IQR method, **{_fmt_number(metrics['high_value_outlier_count_iqr'])} invoices** were identified as high-value outliers, representing **{_fmt_percent(metrics['high_value_outlier_share_percent'])}** of invoices.",
            "",
            "![Invoice Value Distribution](../outputs/figures/04_invoice_value_distribution.png)",
            "",
            "This chart shows the distribution of invoice values up to the 99th percentile so that normal purchase behavior is visible without extreme invoices dominating the chart. The outlier analysis is important because the dataset includes wholesale customers, so very large invoices may represent bulk buying behavior rather than ordinary individual shopping.",
            "",
            "### 6.5 Relationship between basket size and invoice value",
            "",
            f"The NumPy correlation between distinct products in an invoice and invoice value is **{_fmt_decimal(metrics['basket_value_correlation_distinct_items'])}**. The correlation between total units and invoice value is **{_fmt_decimal(metrics['basket_value_correlation_total_units'])}**.",
            "",
            "![Relationship Between Distinct Products and Invoice Value](../outputs/figures/05_invoice_items_vs_value_relationship.png)",
            "",
            "This scatter plot supports the relationship analysis requirement. A positive correlation means invoices containing more products or units tend to have higher values, although correlation does not prove causation.",
            "",
            "### 6.6 UK vs international invoice value comparison",
            "",
            f"The United Kingdom accounts for **{_fmt_percent(metrics['uk_revenue_share_percent'])}** of total revenue in the cleaned dataset.",
            "",
            "![UK vs International Boxplot](../outputs/figures/06_market_group_boxplot.png)",
            "",
            "This boxplot compares invoice values between United Kingdom and international transactions. It is useful because the company is UK-based, so comparing domestic and international purchase behavior helps show whether foreign customers behave differently from the main market.",
            "",
            "## 7. Key findings",
            "",
            f"1. The dataset is large enough for meaningful analysis, with **{_fmt_number(metrics['raw_rows'])} raw transaction rows** and **{_fmt_number(metrics['analysis_ready_rows'])} positive-sales rows** after cleaning.",
            f"2. The business is highly influenced by its strongest market: **{metrics['top_country']}** contributes **{_fmt_percent(metrics['top_country_share_percent'])}** of total revenue.",
            f"3. The strongest revenue month is **{metrics['best_month']}**, which suggests that revenue is not evenly distributed across the year.",
            f"4. Product revenue is concentrated: the top product, **{metrics['top_product']}**, generates **{_fmt_money(metrics['top_product_revenue'])}** in revenue.",
            f"5. Basket size and invoice value have a correlation of **{_fmt_decimal(metrics['basket_value_correlation_distinct_items'])}**, showing that invoices with more distinct products generally tend to be more valuable.",
            f"6. The IQR outlier method found **{_fmt_number(metrics['high_value_outlier_count_iqr'])} high-value invoices**, which is important because wholesale orders can strongly affect averages.",
            "",
            "## 8. Limitations",
            "",
            "- The dataset covers a limited time period, so the results may not represent long-term retail behavior.",
            "- Some customer IDs are missing, so customer-level behavior analysis cannot include every transaction.",
            "- The company is UK-based, so results may be biased toward the United Kingdom market.",
            "- Removing cancellations and returns is suitable for positive sales analysis, but it means the project does not fully analyze refund behavior.",
            "- Correlation analysis shows association between variables, but it does not prove that one variable causes another.",
            "- Product descriptions may contain inconsistencies, which can affect product-level grouping.",
            "",
            "## 9. Conclusion",
            "",
            "This project analyzed online retail transaction data using Pandas, NumPy, and Matplotlib. The analysis found clear revenue trends over time, strong country and product concentration, meaningful customer and basket-level patterns, and a group of unusually high-value invoices. The findings show how transaction data can support business decisions such as market prioritization, inventory planning, seasonal preparation, and customer-value analysis.",
            "",
            "## 10. Files produced by the project",
            "",
            "- Cleaned dataset: `data/processed/cleaned_online_retail.csv`",
            "- Summary tables: `outputs/tables/`",
            "- Matplotlib figures: `outputs/figures/`",
            "- Report: `report/Final_Project_Report.md` and `report/Final_Project_Report.docx`",
            "",
        ]
    )

    report_path.write_text("\n".join(lines), encoding="utf-8")
    return report_path


def generate_docx_report(student_name: str, student_id: str) -> Path:
    """
    Generate a DOCX report.

    This function is intentionally simple and based on the Markdown report content.
    The Markdown report remains the source of truth.
    """
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Run `pip install python-docx` or use the Markdown report."
        ) from exc

    results = _load_json(ANALYSIS_RESULTS_FILE)
    metrics = results["key_metrics"]
    cleaning_log = _load_json(CLEANING_LOG_FILE)

    doc = Document()
    doc.add_heading(PROJECT_TITLE, 0)
    doc.add_paragraph(f"Student name: {student_name}")
    doc.add_paragraph(f"Student ID: {student_id}")
    doc.add_paragraph(f"Dataset: {DATASET_NAME}")
    doc.add_paragraph(f"Source: {DATASET_SOURCE_PAGE}")
    doc.add_paragraph(f"Citation: {DATASET_CITATION}")

    doc.add_heading("1. Project objective", level=1)
    doc.add_paragraph(
        "The objective of this project is to analyze customer purchase behavior and revenue trends "
        "using a real online retail transaction dataset. The analysis focuses on revenue trends, "
        "country and product contribution, customer behavior, and unusual high-value transactions."
    )

    doc.add_heading("2. Analytical questions", level=1)
    for question in [
        "How does online retail revenue change over time?",
        "Which countries and products contribute most to total revenue?",
        "What patterns exist in customer purchase behavior, and which transactions look unusually large?",
    ]:
        doc.add_paragraph(question, style="List Number")

    doc.add_heading("3. Dataset description", level=1)
    doc.add_paragraph(
        f"The raw dataset contains {_fmt_number(metrics['raw_rows'])} rows and "
        f"{_fmt_number(metrics['raw_columns'])} columns. After cleaning, the analysis-ready sales "
        f"dataset contains {_fmt_number(metrics['analysis_ready_rows'])} rows. The transaction period "
        f"is from {metrics['date_min']} to {metrics['date_max']}."
    )

    doc.add_heading("4. Data cleaning and preparation", level=1)
    for item in cleaning_log["cleaning_justification"]:
        doc.add_paragraph(f"{item['step']}: {item['reason']}", style="List Bullet")

    doc.add_heading("5. Feature engineering", level=1)
    features = [
        "total_revenue = quantity × unit_price",
        "invoice_month, invoice_year, invoice_day_name, and invoice_hour",
        "market_group for United Kingdom vs International comparison",
        "day_type for Weekday vs Weekend comparison",
        "basket_value and basket_distinct_items for invoice-level analysis",
        "basket_value_category and line_revenue_zscore using NumPy calculations",
    ]
    for feature in features:
        doc.add_paragraph(feature, style="List Bullet")

    doc.add_heading("6. Analysis and visualizations", level=1)
    doc.add_paragraph(
        f"Total positive sales revenue is {_fmt_money(metrics['total_revenue'])}, "
        f"with {_fmt_number(metrics['total_orders'])} unique invoices and an average order value of "
        f"{_fmt_money(metrics['average_order_value'])}."
    )

    figure_sections = [
        (
            "Monthly revenue trend",
            FIGURES_DIR / "01_monthly_revenue_trend.png",
            f"The highest revenue month is {metrics['best_month']} with revenue of {_fmt_money(metrics['best_month_revenue'])}.",
        ),
        (
            "Top countries by revenue",
            FIGURES_DIR / "02_top_countries_by_revenue.png",
            f"The top country is {metrics['top_country']}, contributing {_fmt_percent(metrics['top_country_share_percent'])} of revenue.",
        ),
        (
            "Top products by revenue",
            FIGURES_DIR / "03_top_products_by_revenue.png",
            f"The top product is {metrics['top_product']} with revenue of {_fmt_money(metrics['top_product_revenue'])}.",
        ),
        (
            "Invoice value distribution",
            FIGURES_DIR / "04_invoice_value_distribution.png",
            f"The IQR method identified {_fmt_number(metrics['high_value_outlier_count_iqr'])} high-value invoice outliers.",
        ),
        (
            "Relationship analysis",
            FIGURES_DIR / "05_invoice_items_vs_value_relationship.png",
            f"The correlation between distinct products and invoice value is {_fmt_decimal(metrics['basket_value_correlation_distinct_items'])}.",
        ),
        (
            "UK vs international comparison",
            FIGURES_DIR / "06_market_group_boxplot.png",
            f"The UK accounts for {_fmt_percent(metrics['uk_revenue_share_percent'])} of cleaned revenue.",
        ),
    ]

    for title, image_path, caption in figure_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(caption)
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(5.8))

    doc.add_heading("7. Key findings", level=1)
    findings = [
        f"The cleaned data contains {_fmt_number(metrics['analysis_ready_rows'])} positive-sales rows.",
        f"{metrics['top_country']} is the strongest revenue country.",
        f"{metrics['best_month']} is the highest revenue month.",
        f"{metrics['top_product']} is the highest revenue product.",
        f"Basket size and invoice value have a correlation of {_fmt_decimal(metrics['basket_value_correlation_distinct_items'])}.",
        f"The IQR method found {_fmt_number(metrics['high_value_outlier_count_iqr'])} high-value invoice outliers.",
    ]
    for finding in findings:
        doc.add_paragraph(finding, style="List Bullet")

    doc.add_heading("8. Limitations", level=1)
    limitations = [
        "The dataset covers a limited time period.",
        "Some customer IDs are missing, limiting customer-level analysis.",
        "The business is UK-based, so results may be biased toward the UK market.",
        "The analysis removes cancellations and returns from positive-sales analysis.",
        "Correlation does not prove causation.",
    ]
    for limitation in limitations:
        doc.add_paragraph(limitation, style="List Bullet")

    doc.add_heading("9. Conclusion", level=1)
    doc.add_paragraph(
        "This project used Pandas, NumPy, and Matplotlib to analyze online retail transactions. "
        "The results show revenue trends, important countries and products, customer purchasing "
        "patterns, and high-value outliers that can support business decisions."
    )

    docx_path = REPORT_DIR / "Final_Project_Report.docx"
    doc.save(docx_path)
    return docx_path
